import streamlit as st
import streamlit.components.v1 as components
import datetime
import io
import os
import base64
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from database import init_db, create_user, verify_user, save_profile, get_all_profiles, delete_profile, save_locked_field, get_locked_fields, delete_locked_field

# Custom HTML5 signature pad (handles touch on mobile + mouse on desktop)
_SIGNATURE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "components", "signature_pad")
_signature_pad = components.declare_component("signature_pad", path=_SIGNATURE_DIR)


def _insert_paragraph_before(paragraph):
    """Create and return a new empty paragraph directly above `paragraph`."""
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def vorlage_befuellen(vorlagen_pfad: str, daten: dict, selected_font: str = "Arial", signature_bytes: bytes = None) -> bytes:
    dok = Document(vorlagen_pfad)

    # Enforce margins
    for section in dok.sections:
        section.top_margin = Mm(35)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(25)
        section.right_margin = Mm(20)

    # Insert the signature image just above the sender's full name, if provided
    if signature_bytes:
        for p in dok.paragraphs:
            if '{{ABSENDER_VOLLNAME}}' in p.text:
                sig_para = _insert_paragraph_before(p)
                # Tight spacing so the signature sits close to the name above/below
                pf = sig_para.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0
                sig_run = sig_para.add_run()
                sig_run.add_picture(io.BytesIO(signature_bytes), width=Mm(38))
                break

    # Optional paragraph removal: remove any paragraph that only contains an empty placeholder
    placeholders_to_check = [
        'EMPFAENGER_FIRMA', 'EMPFAENGER_ABTEILUNG', 'EMPFAENGER_STRASSE', 'EMPFAENGER_PLZ', 'EMPFAENGER_ORT',
        'EINLEITUNGSSATZ', 'HAUPTTEIL_ABSATZ_1', 'HAUPTTEIL_ABSATZ_2', 'HAUPTTEIL_ABSATZ_3'
    ]
    for p in list(dok.paragraphs):
        text = p.text
        for ph in placeholders_to_check:
            ph_token = f"{{{{{ph}}}}}"
            if ph_token in text and not daten.get(ph, '').strip():
                p._element.getparent().remove(p._element)
                break

    def in_absatz_ersetzen(absatz, daten):
        for run in absatz.runs:
            for schluessel, wert in daten.items():
                platzhalter = f"{{{{{schluessel}}}}}"
                if platzhalter in run.text:
                    run.text = run.text.replace(platzhalter, wert)
                    run.font.color.rgb = None
                    run.font.underline = False
            # Override font for all runs
            run.font.name = selected_font

    for absatz in dok.paragraphs:
        in_absatz_ersetzen(absatz, daten)

    puffer = io.BytesIO()
    dok.save(puffer)
    puffer.seek(0)
    return puffer.getvalue()

# Initialize DB once per session (not on every rerun)
if not st.session_state.get("db_ready"):
    try:
        init_db()
        st.session_state["db_ready"] = True
    except Exception as e:
        st.error(
            "Could not connect to the database. Make sure DATABASE_URL is set in "
            "the app's Secrets (Supabase connection string).\n\n"
            f"Details: {e}"
        )
        st.stop()

st.set_page_config(page_title="applyfly", page_icon="📄", layout="wide")

st.markdown("""
<style>
/* --- Responsive layout --- */
@media (max-width: 768px) {
    div[data-testid="stHorizontalBlock"] { flex-wrap: wrap; }
    div[data-testid="stColumn"] { min-width: 100% !important; }
}

/* --- Lock toggle: hide the checkbox square entirely --- */
div[data-testid="stCheckbox"] [data-baseweb="checkbox"] > div:first-child {
    display: none !important;
}
/* Remove extra left padding from hidden box */
div[data-testid="stCheckbox"] label[data-baseweb="checkbox"] {
    padding-left: 0 !important;
    gap: 0 !important;
}
/* Unchecked → grey, small font */
div[data-testid="stCheckbox"] {
    filter: grayscale(100%);
    opacity: 0.35;
    transition: filter 0.15s ease, opacity 0.15s ease;
    cursor: pointer;
    font-size: 1.2rem;
    line-height: 1;
}
/* Checked → full colour */
div[data-testid="stCheckbox"]:has(input:checked) {
    filter: none;
    opacity: 1;
}

/* --- Prevent the UI from fading/dimming during reruns --- */
[data-stale="true"] { opacity: 1 !important; }
.element-container { opacity: 1 !important; transition: none !important; }
.stApp [data-testid="stStatusWidget"] { transition: none !important; }
</style>
""", unsafe_allow_html=True)

# Disable browser autofill/autocorrect on all inputs. Autofill was writing into
# fields, firing reruns (the periodic fade) and occasionally causing errors.
components.html(
    """
    <script>
    const doc = window.parent.document;
    function disableAutofill() {
        doc.querySelectorAll('input, textarea').forEach(function (el) {
            el.setAttribute('autocomplete', el.type === 'password' ? 'new-password' : 'off');
            el.setAttribute('autocorrect', 'off');
            el.setAttribute('autocapitalize', 'off');
            el.setAttribute('spellcheck', 'false');
        });
    }
    disableAutofill();
    new MutationObserver(disableAutofill).observe(doc.body, {childList: true, subtree: true});
    </script>
    """,
    height=0,
)

# Initialize session state variables
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_email" not in st.session_state:
    st.session_state.user_email = ""

def login_flow():
    st.title("Welcome to :red[applyfly]")
    st.subheader("Your Privacy-First German Cover Letter Formatter")
    
    tab1, tab2 = st.tabs(["Login", "Sign Up"])
    
    with tab1:
        st.markdown("### Login")
        login_email = st.text_input("Email", key="login_email")
        login_password = st.text_input("Password", type="password", key="login_password")
        if st.button("Login"):
            login_email_norm = login_email.strip().lower()
            if verify_user(login_email_norm, login_password):
                st.session_state.logged_in = True
                st.session_state.user_email = login_email_norm
                st.success("Logged in successfully!")
                st.rerun()
            else:
                st.error("Invalid email or password.")
                
    with tab2:
        st.markdown("### Sign Up")
        signup_first_name = st.text_input("First Name")
        signup_last_name = st.text_input("Last Name")
        signup_email = st.text_input("Email", key="signup_email")
        signup_password = st.text_input("Password", type="password", key="signup_password")
        signup_password_confirm = st.text_input("Confirm Password", type="password")
        
        if st.button("Sign Up"):
            signup_email_norm = signup_email.strip().lower()
            if signup_password != signup_password_confirm:
                st.error("Passwords do not match.")
            elif not signup_email_norm or not signup_first_name.strip() or not signup_last_name.strip() or not signup_password:
                st.error("Please fill in all fields.")
            elif "@" not in signup_email_norm or "." not in signup_email_norm.split("@")[-1]:
                st.error("Please enter a valid email address.")
            else:
                if create_user(signup_email_norm, signup_first_name.strip(), signup_last_name.strip(), signup_password):
                    st.success("Account created successfully! You can now log in.")
                else:
                    st.error("Email already exists.")

if not st.session_state.logged_in:
    login_flow()
else:
    # Main Dashboard Sidebar
    st.sidebar.title(f"Hello, {st.session_state.user_email.split('@')[0]}!")
    if st.sidebar.button("Logout"):
        # Clear everything so the next user starts clean (no leftover locks/inputs)
        for _k in list(st.session_state.keys()):
            del st.session_state[_k]
        st.session_state.logged_in = False
        st.session_state.user_email = ""
        st.rerun()
        
    st.title(":red[applyfly] Dashboard")

    # Fetch profiles once per run; both tabs reuse this.
    profiles = get_all_profiles(st.session_state.user_email)

    tab_gen, tab_prof = st.tabs(["📄 Generate Document", "💾 Manage Profiles"])

    with tab_prof:
        st.header("Manage Sender Profiles")
        
        # Create new profile
        with st.expander("Create New Profile"):
            new_profile_name = st.text_input("Profile Name (e.g., 'Software Dev Profile')")
            new_name = st.text_input("Full Name")
            new_street = st.text_input("Street Address")
            new_zip = st.text_input("ZIP Code")
            new_city = st.text_input("City")
            new_phone = st.text_input("Phone Number")
            new_email = st.text_input("Email Address")

            st.markdown("**Signature (optional)**")
            st.caption("Press **Draw** to activate the pad, sign with your finger (mobile) or mouse (desktop), then press **OK** to confirm before saving.")
            if "sig_pad_ver" not in st.session_state:
                st.session_state.sig_pad_ver = 0
            sig_dataurl = _signature_pad(key=f"signature_pad_{st.session_state.sig_pad_ver}", default="")
            if sig_dataurl:
                st.success("Signature confirmed — it will be saved with this profile.")

            if st.button("Save Profile"):
                if new_profile_name and new_name and new_street and new_zip and new_city and new_phone and new_email:
                    # The component returns a data URL ("data:image/png;base64,....");
                    # keep just the base64 payload to match the stored format.
                    signature_b64 = ""
                    if sig_dataurl and "," in sig_dataurl:
                        signature_b64 = sig_dataurl.split(",", 1)[1]
                    save_profile(st.session_state.user_email, new_profile_name, new_name, new_street, new_phone, new_email, new_city, new_zip, signature_b64)
                    # Reset the pad so the next new profile starts blank
                    st.session_state.sig_pad_ver += 1
                    st.success("Profile saved successfully!")
                    st.rerun()
                else:
                    st.error("Please fill in all fields.")
                    
        # List and delete profiles
        st.subheader("Saved Profiles")
        if not profiles:
            st.info("No saved profiles found.")
        else:
            for p in profiles:
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.write(f"**{p['profile_name']}** ({p['name']} - {p['email']})")
                    if p.get('signature'):
                        st.image(base64.b64decode(p['signature']), width=160, caption="Signature")
                with col2:
                    if st.button("Delete", key=f"del_{p['id']}"):
                        delete_profile(p['id'], st.session_state.user_email)
                        st.rerun()
                        
    with tab_gen:
        st.header("Generate Cover Letter")
        
        selected_font = st.selectbox("Select Document Font", ["Arial", "Times New Roman"])
        
        # Seed session_state locks from DB only once per login session
        if "locks" not in st.session_state:
            st.session_state.locks = get_locked_fields(st.session_state.user_email)
        locked_fields = st.session_state.locks
        
        def render_locked_input(label, key, default_val="", is_textarea=False, **kwargs):
            col_lock, col_input = st.columns([1, 11])
            locked = key in locked_fields
            input_key = f"input_{key}"

            # When a *locked* field is edited, persist the new value so it
            # survives refresh and logout without re-toggling the lock.
            def _on_input_change(k=key):
                if st.session_state.get(f"check_lock_{k}"):
                    current_val = st.session_state.get(f"input_{k}", default_val)
                    save_locked_field(st.session_state.user_email, k, current_val)
                    st.session_state.locks[k] = current_val

            with col_input:
                # Show locked value if available, otherwise default
                val = locked_fields.get(key, default_val)
                if is_textarea:
                    input_val = st.text_area(label, value=val, key=input_key, on_change=_on_input_change, **kwargs)
                else:
                    input_val = st.text_input(label, value=val, key=input_key, on_change=_on_input_change, **kwargs)

            with col_lock:
                st.write("")
                st.write("")
                # on_change callback fires immediately when toggled — no rerun needed
                def _on_toggle(k=key):
                    new_state = st.session_state[f"check_lock_{k}"]
                    current_val = st.session_state.get(f"input_{k}", default_val)
                    if new_state:
                        save_locked_field(st.session_state.user_email, k, current_val)
                        st.session_state.locks[k] = current_val
                    else:
                        delete_locked_field(st.session_state.user_email, k)
                        st.session_state.locks.pop(k, None)

                st.checkbox(
                    "🔒",
                    value=locked,
                    key=f"check_lock_{key}",
                    on_change=_on_toggle,
                )

            return input_val

        profile_options = {p['profile_name']: p for p in profiles}
        
        selected_profile_name = st.selectbox("Select Sender Profile", options=["-- Select a Profile --"] + list(profile_options.keys()))
        
        selected_profile = profile_options.get(selected_profile_name, {}) if selected_profile_name != "-- Select a Profile --" else {}
        
        st.subheader("Sender Information")
        col1, col2 = st.columns(2)
        with col1:
            sender_name = st.text_input("Sender Name", value=selected_profile.get('name', ''))
            sender_street = st.text_input("Sender Street", value=selected_profile.get('street', ''))
            sender_zip = st.text_input("Sender ZIP Code", value=selected_profile.get('zip_code', ''))
            sender_city = st.text_input("Sender City", value=selected_profile.get('city', ''))
        with col2:
            sender_phone = st.text_input("Sender Phone", value=selected_profile.get('phone', ''))
            sender_email = st.text_input("Sender Email", value=selected_profile.get('email', ''))
            
        st.subheader("Recipient Information")
        recipient_company = render_locked_input("Company / Authority", "recipient_company")
        recipient_department = render_locked_input("Department / PO Box (Optional)", "recipient_department")
        recipient_street = render_locked_input("Recipient Street (leave empty if PO Box)", "recipient_street")
        recipient_zip = render_locked_input("Recipient ZIP Code", "recipient_zip")
        recipient_city = render_locked_input("Recipient City", "recipient_city")
            
        st.subheader("Date & Subject")
        ort = render_locked_input("Location", "ort", default_val=selected_profile.get('city', ''))
        datum = render_locked_input("Date", "datum", default_val=datetime.datetime.now().strftime("%d.%m.%Y"))
        job_title = render_locked_input("Subject", "job_title", default_val="Bewerbung als ... (Kennziffer ...)")
            
        st.subheader("Letter Contents")
        salutation = render_locked_input("Salutation", "salutation", default_val="Sehr geehrte Damen und Herren,", max_chars=100)
        einleitung = render_locked_input("Introduction (Paragraph 1)", "einleitung", is_textarea=True, max_chars=400, placeholder="Max 400 characters to keep it on one page...")
        body = render_locked_input("Cover Letter Body", "body", is_textarea=True, max_chars=2000, height=260, placeholder="Write the full body of your cover letter here. Separate paragraphs with a blank line. Keep it concise to stay on one page.")
        schlusssatz = render_locked_input("Closing Sentence", "schlusssatz", default_val="Ich freue mich auf Ihre Einladung.", max_chars=200, placeholder="Max 200 characters...")
        grussformel = render_locked_input("Valediction", "grussformel", default_val="Mit freundlichen Grüßen,", max_chars=100)
        
        if st.button("Generate Document"):
            template_path = os.path.join('cover', 'din5008_bewerbung_vorlage.docx')
            if not os.path.exists(template_path):
                st.error(f"Template '{template_path}' not found.")
            else:
                try:
                    # Split the single body into up to 3 paragraphs (blank line =
                    # paragraph break) to fill the template's body placeholders.
                    body_parts = [p.strip() for p in body.split("\n\n") if p.strip()]
                    absatz1 = body_parts[0] if len(body_parts) > 0 else ""
                    absatz2 = body_parts[1] if len(body_parts) > 1 else ""
                    absatz3 = "\n".join(body_parts[2:]) if len(body_parts) > 2 else ""

                    daten = {
                        "ABSENDER_NAME": sender_name,
                        "ABSENDER_STRASSE": sender_street,
                        "ABSENDER_PLZ": sender_zip,
                        "ABSENDER_ORT": sender_city,
                        "ABSENDER_TELEFON": sender_phone,
                        "ABSENDER_EMAIL": sender_email,
                        "EMPFAENGER_FIRMA": recipient_company,
                        "EMPFAENGER_ABTEILUNG": recipient_department,
                        "EMPFAENGER_STRASSE": recipient_street,
                        "EMPFAENGER_PLZ": recipient_zip,
                        "EMPFAENGER_ORT": recipient_city,
                        "ORT": ort,
                        "DATUM": datum,
                        "BETREFF": job_title + "\n",
                        "ANREDE": salutation,
                        "EINLEITUNGSSATZ": einleitung,
                        "HAUPTTEIL_ABSATZ_1": absatz1,
                        "HAUPTTEIL_ABSATZ_2": absatz2,
                        "HAUPTTEIL_ABSATZ_3": absatz3,
                        "SCHLUSSSATZ": schlusssatz,
                        "GRUSSFORMEL": grussformel,
                        "ABSENDER_VOLLNAME": sender_name,
                    }
                    
                    # Decode the selected profile's signature, if it has one
                    signature_bytes = None
                    if selected_profile.get('signature'):
                        signature_bytes = base64.b64decode(selected_profile['signature'])

                    docx_bytes = vorlage_befuellen(template_path, daten, selected_font, signature_bytes)
                    
                    st.success("Document generated successfully!")
                    st.download_button(
                        label="Download Cover Letter (.docx)",
                        data=docx_bytes,
                        file_name=f"Cover_Letter_{sender_name.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Error generating document: {e}")
